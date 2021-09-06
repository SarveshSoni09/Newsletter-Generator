from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
from docx.shared import Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import numpy as np

# Red Shade: #A93639
# Blue Shade: #005BAA

def head_of_newsletter(p):
    year = input("Enter Year of Magazine: (Format: 20XX-20XX)")
    sem = input("Enter Semester (Odd/Even): ")
    vol = input("Enter Volume: ")

    p2 = "Academic Year " + year
    p4 = "For " + sem + " Sem, Volume " + vol

    head1(p[2], p2)
    head2(p[4], p4)


def head1(para, content):
    font_name = 'Footlight MT Light'
    font_size = Pt(18)
    font_color = RGBColor(0xA9, 0x36, 0x39)
    font_bold = True

    run = para.runs
    r = run[0]
    r.text = content
    font = r.font

    font.name = font_name
    font.size = font_size
    font.bold = font_bold
    font.color.rgb = font_color


def head2(para, content):
    font_name = 'Footlight MT Light'
    font_size = Pt(14)
    font_color = RGBColor(0xA9, 0x36, 0x39)
    font_bold = True

    run = para.runs
    r = run[0]
    r.text = content
    font = r.font

    font.name = font_name
    font.size = font_size
    font.bold = font_bold
    font.color.rgb = font_color

def group_pic():
    para = p[5]
    para.add_run("")
    run = para.runs
    r = run[0]
    img_name = input("Enter name of Image (xyz.png/xyz.jpg): ")
    r.add_picture(img_name, width=Inches(7))

def blue_head(r):
    font_name = 'Footlight MT Light'
    font_size = Pt(12)
    font_color = RGBColor(0x00, 0x5B, 0xAA)
    font_bold = True

    # run = para.runs
    # r = run[0]
    # r.text = content
    font = r.font

    font.name = font_name
    font.size = font_size
    font.bold = font_bold
    font.color.rgb = font_color


def red_head(r):
    font_name = 'Footlight MT Light'
    font_size = Pt(12)
    font_color = RGBColor(0xA9, 0x36, 0x39)
    font_bold = True

    # run = para.runs
    # r = run[0]
    # r.text = content
    font = r.font

    font.name = font_name
    font.size = font_size
    font.bold = font_bold
    font.color.rgb = font_color


def red_content(r):
    font_name = 'Footlight MT Light'
    font_size = Pt(12)
    font_color = RGBColor(0xA9, 0x36, 0x39)
    font_bold = False

    # run = para.runs
    # r = run[0]
    # r.text = content
    font = r.font

    font.name = font_name
    font.size = font_size
    font.bold = font_bold
    font.color.rgb = font_color


document = Document('Newsletter Template.docx')



# p[5].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER



# head_of_newsletter(p)
# group_pic()

# headings = ["Highlights of the Department", "Remarkable Milestones"]

hlts_of_dept = ["Lorem ipsum dolor sit amet, consectetur adipiscing elit. Curabitur pulvinar dignissim ex.",
                "Donec nunc leo, accumsan pretium nisl ac, bibendum fringilla ligula.",
                "Aliquam erat volutpat. In eget facilisis orci. Aenean iaculis euismod nulla, nec convallis diam ornare volutpat."]

rem_mlstns = ["Curabitur venenatis, nibh quis pharetra tincidunt, odio nisl sollicitudin justo, vel efficitur odio urna eu libero.",
              "Praesent id urna arcu. Curabitur ultricies enim ac sapien dictum, non lacinia metus tincidunt. Nulla tellus urna.",
              "facilisis sit amet augue eu, semper imperdiet nisi. Nulla quis porta tellus, sit amet ornare nunc."]

# fac_ach = []
# stud_ach = []

head_content = {"Highlights of the Department": hlts_of_dept,
                "Remarkable Milestones": rem_mlstns}

# for content in head_content:
#     document.add_paragraph()
#     p = list(document.paragraphs)
#     p[-1].paragraph_format.left_indent = Inches(-0.3333)
#     p[-1].paragraph_format.space_before = Pt(10)
#     p[-1].paragraph_format.space_after = Pt(2)
#     p[-1].add_run(content)
#     blue_head(p[-1].runs[0])
#     sub_content = head_content.get(content)
#     for cont in sub_content:
#         document.add_paragraph()
#         p = list(document.paragraphs)
#         p[-1].paragraph_format.left_indent = Inches(-0.3333)
#         p[-1].paragraph_format.space_before = Pt(0)
#         p[-1].paragraph_format.space_after = Pt(2)
#         p[-1].add_run().add_picture('bullet_red.png', width=Inches(0.06))
#         p[-1].add_run("   "+cont)
#         red_content(p[-1].runs[1])

table_content = [
    ['Heading of Pic 1', 'CrispyChillyPotatoes.jpg', 'Caption of Pic 1'],
    ['Heading of Pic 2', 'GarlicBallSizzler.jpg', 'Caption of Pic 2'],
    ['Heading of Pic 3', 'PaneerTikkaSizzler.jpg', 'Caption of Pic 3'],
    ['Heading of Pic 4', 'Quesadilla.jpg', 'Caption of Pic 4']
]

table_content = np.array(table_content)
head_of_pics = []
img_names = []
caption_of_pics = []
# for content in table_content:
#     print(content)

# for i in range(len(table_content)):
#     pass
for head_of_pic, img_name, caption_of_pic in table_content:
    head_of_pics.append(head_of_pic)
    img_names.append(img_name)
    caption_of_pics.append(caption_of_pic)

head_of_pics = np.array(head_of_pics)
img_names = np.array(img_names)
caption_of_pics = np.array(caption_of_pics)

head_of_pics = np.reshape(head_of_pics, (-1, 2))
img_names = np.reshape(img_names, (-1, 2))
caption_of_pics = np.reshape(caption_of_pics, (-1, 2))

new_table_of_content = []
table = document.add_table(rows=1, cols=2)
for i in range(len(head_of_pics)):
    # new_table_of_content.append(head_of_pics[i])
    row_cells = table.add_row().cells
    paragraph = row_cells[0].paragraphs[0]
    run = paragraph.add_run(head_of_pics[i][0])
    blue_head(run)
    paragraph = row_cells[1].paragraphs[0]
    run = paragraph.add_run(head_of_pics[i][1])
    blue_head(run)
    # new_table_of_content.append(img_names[i])
    row_cells = table.add_row().cells
    paragraph = row_cells[0].paragraphs[0]
    run = paragraph.add_run()
    run.add_picture(img_names[i][0], width=Inches(3), height=Inches(2))
    paragraph = row_cells[1].paragraphs[0]
    run = paragraph.add_run()
    run.add_picture(img_names[i][1], width=Inches(3), height=Inches(2))
    # new_table_of_content.append(caption_of_pics[i])
    row_cells = table.add_row().cells
    paragraph = row_cells[0].paragraphs[0]
    run = paragraph.add_run(caption_of_pics[i][0])
    red_head(run)
    paragraph = row_cells[1].paragraphs[0]
    run = paragraph.add_run(caption_of_pics[i][1])
    red_head(run)


# table = document.add_table(rows=1, cols=2)
# hdr_cells = table.rows[0].cells
# for content in new_table_of_content:
#     row_cells = table.add_row().cells
#     row_cells[0].text = content[0]
#     paragraph = row_cells[1].paragraphs[0]
#     run = paragraph.add_run()
#     # run.add_picture(img_name, width=Inches(3), height=Inches(2))
#     # row_cells[1].add_picture(img_name, width=Inches(4), height=Inches(3))
#     row_cells[1].text = content[1]


document.save('Newsletter Template.docx')
# p = list(document.paragraphs)
# for line, para in enumerate(p):
#     print("Paragraph {}:".format(line))
#     print(para.text)
#     print()

# document.save('Newsletter Template.docx')
