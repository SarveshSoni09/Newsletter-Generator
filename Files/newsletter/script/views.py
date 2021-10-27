import mimetypes
import os

import lxml
import matplotlib.pyplot as plt
import numpy as np
from admin_panel import views as admv
from admin_panel.models import Header
from django.conf import settings
from django.http import Http404, HttpResponse
from django.http.response import HttpResponse
from django.shortcuts import render
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Inches, Pt, RGBColor
from faculty_panel.models import *
from newsletter.settings import BASE_DIR
from newsletter.views import *


# Red Shade: #A93639
# font_color = RGBColor(0xA9, 0x36, 0x39)
# Blue Shade: #005BAA
# font_color = RGBColor(0x00, 0x5B, 0xAA)
# Purple Shade: #7864A2
# font_color = RGBColor(0x78, 0x64, 0xA2)
i=1

def newsletter(request):
    if request.method == "GET" and request.user.is_authenticated:
        def para_format(p, lin, rin, spb, spa):
            p.paragraph_format.left_indent = Inches(lin)
            p.paragraph_format.right_indent = Inches(rin)
            p.paragraph_format.space_before = Pt(spb)
            p.paragraph_format.space_after = Pt(spa)


        def head1(r):
            font = r.font
            font.name = 'Footlight MT Light'
            font.size = Pt(18)
            font.bold = True
            font.color.rgb = RGBColor(0xA9, 0x36, 0x39)


        def head2(r):
            font = r.font
            font.name = 'Footlight MT Light'
            font.size = Pt(24)
            font.bold = True
            font.color.rgb = RGBColor(0x78, 0x64, 0xA2)


        def head3(r):
            font = r.font
            font.name = 'Footlight MT Light'
            font.size = Pt(14)
            font.bold = True
            font.color.rgb = RGBColor(0xA9, 0x36, 0x39)


        def head_of_newsletter(h1, h2, h3, h4):
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = paragraph.add_run(h1)
            head1(run)

            paragraph = document.add_paragraph()
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = paragraph.add_run(h2)
            head1(run)

            paragraph = document.add_paragraph()
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = paragraph.add_run(h3)
            head2(run)

            paragraph = document.add_paragraph()
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = paragraph.add_run(h4)
            head3(run)


        def group_pic(picname):
            para = document.add_paragraph()
            para.paragraph_format.left_indent = Inches(-0.3333)
            para.add_run("")
            run = para.runs
            r = run[0]
            img_name = picname
            r.add_picture(img_name, width=Inches(7))


        def static_content(abt_department, vision, missions, peos):
            paragraph = document.add_paragraph()
            para_format(paragraph, -0.3333, -0.3333, 0, 0)
            run = paragraph.add_run("About Department:")
            blue_head.apply_style(run)

            paragraph = document.add_paragraph()
            para_format(paragraph, -0.3333, -0.3333, 0, 0)
            run = paragraph.add_run(abt_department)
            red_content.apply_style(run)

            paragraph = document.add_paragraph()
            para_format(paragraph, -0.3333, -0.3333, 4, 0)
            run = paragraph.add_run("Vision:")
            blue_head.apply_style(run)

            paragraph = document.add_paragraph()
            para_format(paragraph, -0.3333, -0.3333, 2, 0)
            run = paragraph.add_run(vision)
            red_content.apply_style(run)
            run.italic = True

            paragraph = document.add_paragraph()
            para_format(paragraph, -0.3333, -0.3333, 4, 0)
            run = paragraph.add_run("Mission:")
            blue_head.apply_style(run)

            for mission in missions:
                paragraph = document.add_paragraph()
                para_format(paragraph, -0.3333, -0.3333, 2, 0)
                run = paragraph.add_run()
                run.add_picture(os.path.join(BASE_DIR, 'newsletter/Static/ScriptStatic/bullet_red.png'), width=Cm(0.2), height=Cm(0.2))
                run = paragraph.add_run("   " + mission)
                red_content.apply_style(run)
                run.italic = True

            paragraph = document.add_paragraph()
            para_format(paragraph, -0.3333, -0.3333, 4, 0)
            run = paragraph.add_run("Program Educational Objectives (PEOs):")
            blue_head.apply_style(run)

            paragraph = document.add_paragraph()
            para_format(paragraph, -0.3333, -0.3333, 0, 0)
            run = paragraph.add_run("Students will possess:")
            red_content.apply_style(run)

            for i in range(len(peos)):
                paragraph = document.add_paragraph()
                para_format(paragraph, -0.3333, -0.3333, 2, 0)
                run = paragraph.add_run("PEO " + str(i + 1) + ": ")
                red_content.apply_style(run)
                run = paragraph.add_run(peos[i])
                red_content.apply_style(run)
                run.italic = True

            document.add_page_break()


        class FontStyle:
            def __init__(self, size, bold, color):
                self.name = 'Footlight MT Light'
                self.size = Pt(size)
                self.bold = bold
                self.color = color

            def apply_style(self, run):
                font = run.font
                font.name = self.name
                font.size = self.size
                font.bold = self.bold
                font.color.rgb = self.color


        blue_head = FontStyle(12, True, RGBColor(0x00, 0x5B, 0xAA))
        red_head = FontStyle(12, True, RGBColor(0xA9, 0x36, 0x39))
        red_content = FontStyle(12, False, RGBColor(0xA9, 0x36, 0x39))


        class BigContent:

            def __init__(self, heading, description, img, img_cap):
                self.heading = heading
                self.desc = description
                self.img = img
                self.img_cap = img_cap

            def create_docx(self):
                if self.desc:
                    paragraph = document.add_paragraph()
                    para_format(paragraph, -0.3333, -0.3333, 10, 2)
                    run = paragraph.add_run(self.heading)
                    blue_head.apply_style(run)

                for i in range(len(self.desc)):
                    # paragraph = document.add_paragraph()
                    # para_format(paragraph, -0.3333, -0.3333, 4, 0)
                    # run = paragraph.add_run(str(i + 1) + ". " + self.titles[i])
                    # red_head.apply_style(run)

                    if self.desc[i]:
                        paragraph = document.add_paragraph()
                        para_format(paragraph, -0.3333, -0.3333, 2, 0)
                        run = paragraph.add_run(self.desc[i])
                        red_content.apply_style(run)

                    # if self.teams[i][0]:
                    #     paragraph = document.add_paragraph()
                    #     para_format(paragraph, 0, -0.3333, 4, 2)
                    #     run = paragraph.add_run("Team Members: ")
                    #     red_head.apply_style(run)

                    #     table = document.add_table(rows=1, cols=2)
                    #     for j in range(len(self.teams[i])):
                    #         if j % 2 == 0:
                    #             if j > 1:
                    #                 row_cells = table.add_row().cells
                    #             else:
                    #                 row_cells = table.rows[0].cells
                    #             paragraph = row_cells[0].paragraphs[0]
                    #             run = paragraph.add_run(str(j + 1) + ". " + self.teams[i][j])
                    #             red_content.apply_style(run)
                    #         else:
                    #             paragraph = row_cells[1].paragraphs[0]
                    #             run = paragraph.add_run(str(j + 1) + ". " + self.teams[i][j])
                    #             red_content.apply_style(run)

                    if self.img[i]:
                        
                        paragraph = document.add_paragraph()
                        para_format(paragraph, -0.3333, -0.3333, 2, 0)
                        run = paragraph.add_run()
                        run.add_picture(self.img[i], width=Inches(6), height=Inches(3))

                    # if self.img_cap[i]:
                    #     paragraph = document.add_paragraph()
                    #     para_format(paragraph, -0.3333, -0.3333, 2, 4)
                    #     run = paragraph.add_run(self.img_cap[i])
                    #     red_content.apply_style(run)


        class SmallContent:
            def __init__(self, heading, title, desc, pics):
                self.heading = heading
                self.title = title
                self.desc = desc
                self.pics = pics


            def create_docx(self):
                updated_pics = []
                if self.pics:
                    for pic in self.pics:
                        if pic:
                            updated_pics.append(pic)

                if self.desc:
                    paragraph = document.add_paragraph()
                    para_format(paragraph, -0.3333, -0.3333, 10, 2)
                    run = paragraph.add_run(self.heading)
                    blue_head.apply_style(run)  

                for i in range(len(self.desc)):
                    paragraph = document.add_paragraph()
                    para_format(paragraph, -0.3333, -0.3333, 4, 0)
                    if self.title:
                        run  = paragraph.add_run()
                        run.add_picture(os.path.join(BASE_DIR, 'newsletter/Static/ScriptStatic/bullet_red.png'), width=Cm(0.2), height=Cm(0.2))
                        if self.title[i]:
                            run = paragraph.add_run('  '+self.title[i] + ': ')
                        else:
                            run = paragraph.add_run('  ')
                        red_head.apply_style(run)
                    else:
                        run  = paragraph.add_run()
                        run.add_picture(os.path.join(BASE_DIR, 'newsletter/Static/ScriptStatic/bullet_red.png'), width=Cm(0.2), height=Cm(0.2))
                    run = paragraph.add_run(self.desc[i])
                    # if self.desc[i]:
                    #     red_head.apply_style(run)
                    #     run = paragraph.add_run(": " + self.desc[i])
                    #     red_content.apply_style(run)
                    # else:
                    red_content.apply_style(run)

                if updated_pics:
                    if len(updated_pics) < 2:
                        paragraph = document.add_paragraph()
                        run = paragraph.add_run()
                        run.add_picture(updated_pics[0], width=Inches(6), height=Inches(3))
                    elif len(updated_pics) < 3:
                        paragraph = document.add_paragraph()
                        for i in range(len(updated_pics)):
                            if updated_pics[i]:
                                run = paragraph.add_run()
                                run.add_picture(updated_pics[i], width=Inches(3), height=Inches(2))
                                paragraph.add_run("  ")
                    elif len(updated_pics) < 4:
                        paragraph = document.add_paragraph()
                        for i in range(len(updated_pics)):
                            if updated_pics[i]:
                                run = paragraph.add_run()
                                run.add_picture(updated_pics[i], width=Inches(2), height=Inches(1.5))
                                paragraph.add_run("  ")
                    else:
                        table = document.add_table(rows=1, cols=2)
                        for j in range(len(updated_pics)):
                            if updated_pics[j]:
                                if j % 2 == 0:
                                    if j > 1:
                                        row_cells = table.add_row().cells
                                    else:
                                        row_cells = table.rows[0].cells
                                    paragraph = row_cells[0].paragraphs[0]
                                    run = paragraph.add_run()
                                    run.add_picture(updated_pics[j], width=Inches(3), height=Inches(2))
                                else:
                                    paragraph = row_cells[1].paragraphs[0]
                                    run = paragraph.add_run()
                                    run.add_picture(updated_pics[j], width=Inches(3), height=Inches(2))


        def chart_gen(title, xlist, ylist, xlabel, ylabel):
            global i
            plt.clf()
            plt.rcParams.update({'font.size': 12})
            plt.bar(xlist, ylist, color='maroon', width=0.6)
            plt.xlabel(xlabel)
            plt.ylabel(ylabel)
            plt.title(title)
            chart_name = 'chart' + str(i) + '.png'
            plt.savefig(chart_name)
            i += 1
            return chart_name


        def table_create(head_of_pics, table_pics, caption_of_pics):
            head_of_pics = np.array(head_of_pics)
            table_pics = np.array(table_pics)
            caption_of_pics = np.array(caption_of_pics)

            head_of_pics = np.reshape(head_of_pics, (-1, 2))
            table_pics = np.reshape(table_pics, (-1, 2))
            caption_of_pics = np.reshape(caption_of_pics, (-1, 2))

            table = document.add_table(rows=1, cols=2)

            mytab = table
            nsmap = mytab._element[0].nsmap  # For namespaces
            searchtag = '{%s}tblPr' % nsmap['w']  # w:tblPr
            mytag = '{%s}tblInd' % nsmap['w']  # w:tblInd
            myw = '{%s}w' % nsmap['w']  # w:w
            mytype = '{%s}type' % nsmap['w']  # w:type
            for elt in mytab._element:
                if elt.tag == searchtag:
                    myelt = lxml.etree.Element(mytag)
                    myelt.set(myw, '-500')
                    myelt.set(mytype, 'dxa')
                    myelt = elt.append(myelt)

            for i in range(len(head_of_pics)):
                row_cells = table.add_row().cells
                paragraph = row_cells[0].paragraphs[0]
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = paragraph.add_run(head_of_pics[i][0])
                blue_head.apply_style(run)
                paragraph = row_cells[1].paragraphs[0]
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = paragraph.add_run(head_of_pics[i][1])
                blue_head.apply_style(run)

                row_cells = table.add_row().cells
                paragraph = row_cells[0].paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(table_pics[i][0], width=Inches(3.6), height=Inches(2.5))
                paragraph = row_cells[1].paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(table_pics[i][1], width=Inches(3.6), height=Inches(2.5))

                row_cells = table.add_row().cells
                paragraph = row_cells[0].paragraphs[0]
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = paragraph.add_run(caption_of_pics[i][0])
                red_head.apply_style(run)
                paragraph = row_cells[1].paragraphs[0]
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = paragraph.add_run(caption_of_pics[i][1])
                red_head.apply_style(run)

        # NewsLetter Front Page Content
        header = Header.objects.latest('id')
        year = getattr(header, 'academic_year')
        sem = getattr(header, 'semester')
        vol = getattr(header,'volume')

        h1 = 'Department of Computer Engineering'
        h2 = 'Academic Year ' + year
        h3 = '"Abhiyanta" Newsletter'
        h4 = 'For ' + sem + ' Semester, Volume ' + str(vol)

        abt_department = 'The Computer Engineering department was established in the year 2001 for providing excellent education to all students with an aim to foster creativity in Learning, research and building a knowledge base along with imparting quality education. The department possesses well qualified and motivated faculty members and support staff. The laboratories are adequately equipped with state-of-the-art facilities. The department offers add-on courses to cope with the latest trends in technology. It organizes industry visits and many series of expert lectures for the students on the latest developing trends in Computer Industry. Students of our department have won prizes in various national and international level paper presentation competition and project exhibitions. Every year a remarkable number of students place in companies such as Infosys, TCS, Accenture, Cognizant, L&T Infotech, CSC, Tech Mahindra, Mastek, TPG Axon, Artistivity etc.'

        vision = 'To be the finest centre of learning by imparting quality teaching in the field of computer engineering.'

        missions = ['To inculcate in student the ability to analyse, design and develop software projects.',
                    'To prepare students to be ready for ever changing scenario in the field of Computer Engineering.',
                    'To help the students to attain and utilize their potential for successful carrier and to the need of society and industry.']

        peos = [
            'Sound knowledge of mathematics, science and engineering fundamentals required for a lifelong career in computer engineering domains.',
            'Skill to design and develop software projects for solving real world problems with the use of modern tools and techniques.',
            'Professional ethics, leadership qualities and social awareness in the students.']

        # NewLetter Content
        # Highlights of Department
        hlts_heading = "Highlights of the Department"
        highlights = get_highlights(request, Highlights.objects.values())
        hlts_faculty = highlights[0]
        hlts_content = highlights[1]
        hlts_pics = highlights[2]

        # Remarkable Milestones
        milestones = get_milestones(request, Milestones.objects.values())
        rem_mlstns_heading = "Remarkable Milestones:"
        rem_mlstns_desc = milestones[0]
        rem_mlstns_pic = milestones[1]


        # Activities Conducted
        activities = get_activities(request, Activities.objects.values())
        activity_heading = 'Activities Conducted for Students'
        activity_desc = activities[0]
        activity_img = activities[1]
        # activity_img_cap = []

        # Placement Statistics
        placements = get_placements(request, Placements.objects.values())
        company_name = placements[0]
        no_of_recruits = placements[1]
        placement_x = "Name of Commpany"
        placement_y = "Number of Students Hired"
        placement_title = "Placement Statistics"
        placement_caption = "Students placed in various Companies"

        # Result Statisctics
        results = get_results(request, Results.objects.values())
        student_year = results[0]
        students_pass = results[1]
        results_x = "Year"
        results_y = "Number of Students Passed"
        results_title = "Result Statistics"
        results_caption = "Percentage of Students Passed"


        # Student Achievements
        achievements = get_students(request, Students.objects.values())
        stud_ach_heading = "Student Achievements"
        stud_ach_desc = achievements[0]
        stud_ach_img = achievements[1]

        # Faculty Events
        events = get_events(request, Events.objects.values())
        event_heading = 'Events conducted by Faculty'
        event_desc = events[0]
        event_img = events[1]

        # Companny Projects
        projects = get_projects(request, Projects.objects.values())
        company_project_heading = "Project Sponsored by Companies"
        project_desc = projects
        project_pics = []

        # PhD Pursuing Faculties
        phds = get_phds(request, Phd.objects.values())
        phd_heading = "PhD Pursuing Faculties"
        phd_desc = phds

        ### Main Code
        document = Document(os.path.join(BASE_DIR, 'newsletter/static/ScriptStatic/Newsletter Template.docx'))
        


        head_of_newsletter(h1, h2, h3, h4)
        # group_pic("group_pic.png")
        static_content(abt_department, vision, missions, peos)

        highlights = SmallContent(hlts_heading, hlts_faculty, hlts_content, hlts_pics)
        highlights.create_docx()

        remarkable_milestones = BigContent(rem_mlstns_heading, rem_mlstns_desc, rem_mlstns_pic, None)
        remarkable_milestones.create_docx()

        activities_conducted = SmallContent(activity_heading, None, activity_desc, activity_img)
        activities_conducted.create_docx()


        table_pics = []
        table_pics.append(chart_gen(placement_title, company_name, no_of_recruits, placement_x, placement_y))
        table_pics.append(chart_gen(results_title, student_year, students_pass, results_x, results_y))

        head_of_pics = [placement_title, results_title]
        caption_of_pics = [placement_caption, results_caption]

        table_create(head_of_pics, table_pics, caption_of_pics)

        student_achievements = BigContent(stud_ach_heading, stud_ach_desc, stud_ach_img, None)
        student_achievements.create_docx()

        faculty_events = SmallContent(event_heading, None, event_desc, event_img)
        faculty_events.create_docx()

        company_projects = SmallContent(company_project_heading, None, project_desc, None)
        company_projects.create_docx()

        phd_faculties = SmallContent(phd_heading, None, phd_desc, None)
        phd_faculties.create_docx()

        document.save('Newsletter.docx')
        global i
        i=1
        
        
        if 'word' in request.GET:
            pdf = False
            response = download_file(request, pdf)
            return response
        if 'pdf' in request.GET:
            pdf = True
            convert(request)
            response = download_file(request, pdf)
            return response

        if 'reset' in request.GET:
            Highlights.objects.all().delete()
            Milestones.objects.all().delete()
            Activities.objects.all().delete()
            Placements.objects.all().delete()
            Results.objects.all().delete()
            Students.objects.all().delete()
            Events.objects.all().delete()
            Projects.objects.all().delete()
            Phd.objects.all().delete()
            auth.logout(request)
            return redirect('/')


        all_data = admv.get_data()
        # download_doc(request, '/Newsletter.docx')
        return render(request, 'admin-panel.html', {'all_data' : all_data})
    else:
        return render(request, 'index.html')

def download_file(request, pdf):
    # from django.utils.encoding import smart_str
    # response = HttpResponse(content_type='application/force-download')
    # response['Content-Disposition'] = 'attachment; filename=%s' % smart_str('Newsletter.docx')
    # response['X-Sendfile'] = smart_str(os.path.join(BASE_DIR, 'newsletter/Newsletter.docx'))
    # return response
    from django.views.static import serve
    if pdf:
        filepath = os.path.join(BASE_DIR, 'Newsletter.pdf')
    else:
        filepath = os.path.join(BASE_DIR, 'Newsletter.docx')
    return serve(request, os.path.basename(filepath), os.path.dirname(filepath))

def convert(request):
    import comtypes.client
    import sys

    wdFormatPDF = 17

    in_file = os.path.join(BASE_DIR, 'Newsletter.docx')
    out_file = os.path.join(BASE_DIR, 'Newsletter.pdf')

    word = comtypes.client.CreateObject('Word.Application')
    doc = word.Documents.Open(in_file)
    doc.SaveAs(out_file, FileFormat=wdFormatPDF)
    doc.Close()
    word.Quit()
